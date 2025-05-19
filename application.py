# application.py (formerly app.py)
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import io
import os

# Rename the Flask app instance to 'application'
application = Flask(__name__) # CHANGED: app -> application

# --- Configuration for logos ---
# These paths are relative to the application.py file
KEMPER_LOGO_PATH = os.path.join(application.root_path, 'static', 'kemper_markt_logo.png') # CHANGED: app -> application
AKTION_LOGO_PATH = os.path.join(application.root_path, 'static', 'sonder_angebot_logo.png') # CHANGED: app -> application
BIO_LOGO_PATH = os.path.join(application.root_path, 'static', 'bio_logo.png') # CHANGED: app -> application

# Define a consistent logo width for all cases
CONSISTENT_LOGO_WIDTH = Inches(1.0)

# --- Route for the main form ---
@application.route('/', methods=['GET', 'POST']) # CHANGED: app -> application
def index():
    if request.method == 'POST':
        product_type = request.form['product_type']
        department = request.form['department']
        manufacturer = request.form.get('manufacturer', '')
        product_name = request.form['product_name']
        sub_product_name = request.form.get('sub_product_name', '')
        has_varieties = request.form.get('has_varieties') == 'true'
        additional_info = request.form.get('additional_info', '')

        quantity_per_pack_str = request.form['quantity_per_pack']
        unit = request.form['unit']
        price = float(request.form['price'])

        deposit = float(request.form.get('deposit', 0.0))
        packaging_type = request.form.get('packaging_type', '')
        is_bio = request.form.get('is_bio') == 'true'

        try:
            if '/' in quantity_per_pack_str:
                calc_quantity = float(quantity_per_pack_str.split('/')[0])
            else:
                calc_quantity = float(quantity_per_pack_str)
        except ValueError:
            calc_quantity = 0.0

        price_per_unit_text = ""
        # Handle unit price calculation for g, ml, L
        if '/' in quantity_per_pack_str and (unit == "g" or unit == "ml" or unit == "L"):
            quantities = [float(q) for q in quantity_per_pack_str.split('/')]
            prices_per_unit = []
            for q in quantities:
                if q == 0:
                    prices_per_unit.append(0.0)
                    continue
                if unit == "g":
                    prices_per_unit.append((price / q) * 1000)
                elif unit == "ml":
                    prices_per_unit.append((price / q) * 1000)
                elif unit == "L": # Calculate price per liter
                    prices_per_unit.append(price / q) 

            # Format prices with comma as decimal separator
            formatted_prices = [f"{p:,.2f}".replace('.', ',') for p in prices_per_unit]
            price_per_unit_text = "/".join(formatted_prices) + "€" # Append € after all prices
            if unit == "g":
                price_per_unit_text = f"1kg={price_per_unit_text}"
            elif unit == "ml":
                price_per_unit_text = f"1L={price_per_unit_text}"
            elif unit == "L":
                price_per_unit_text = f"1L={price_per_unit_text}" # For L, it's already 1L=...

        elif unit == "g":
            if calc_quantity > 0:
                price_per_kilo = (price / calc_quantity) * 1000
                price_per_unit_text = f"1kg={price_per_kilo:,.2f}€".replace('.', ',')
        elif unit == "ml":
            if calc_quantity > 0:
                price_per_liter = (price / calc_quantity) * 1000
                price_per_unit_text = f"1l={price_per_liter:,.2f}€".replace('.', ',')
        elif unit == "L": # Handle L for single quantities
            if calc_quantity > 0:
                price_per_liter = price / calc_quantity
                price_per_unit_text = f"1L={price_per_liter:,.2f}€".replace('.', ',')
        elif unit == "stück":
            if calc_quantity > 0:
                price_per_piece = price / calc_quantity
                price_per_unit_text = f"1Stück={price_per_piece:,.2f}€".replace('.', ',')


        logo_to_use = KEMPER_LOGO_PATH
        if product_type == 'Aktion':
            logo_to_use = AKTION_LOGO_PATH
        elif department == 'Obst&Gemüse' and is_bio:
            logo_to_use = BIO_LOGO_PATH

        document = generate_document(
            logo_to_use,
            department,
            product_type,
            manufacturer,
            product_name,
            sub_product_name,
            has_varieties,
            additional_info,
            quantity_per_pack_str,
            unit,
            price,
            deposit,
            packaging_type,
            price_per_unit_text,
            is_bio
        )

        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        # Construct the new filename: Marke + Produkt + Unit.docx
        filename_parts = []
        if manufacturer:
            filename_parts.append(manufacturer)
        filename_parts.append(product_name)
        filename_parts.append(f"{quantity_per_pack_str}{unit}")

        # Join parts with a space, then clean up for filename safety
        base_name = " ".join(filter(None, filename_parts)).strip()
        # Replace characters that might cause issues in filenames, like '/'
        base_name = base_name.replace('/', '-')

        filename = f"{base_name}.docx"

        return send_file(doc_io, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return render_template('index.html')

# --- Helper function to add a paragraph with specific style and tight spacing ---
def add_tight_paragraph(doc, text, font_name='Calibri', font_size=None, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0, apply_highlight=False, underline=False, font_color=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if underline:
        run.underline = True
    if font_color:
        run.font.color.rgb = font_color # Set font color

    paragraph.alignment = alignment

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(spacing_after)

    if apply_highlight:
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    return paragraph

# --- Function to generate the Word document ---
def generate_document(logo_path, department, product_type, manufacturer, product_name, sub_product_name,
                    has_varieties, additional_info, quantity_per_pack_str, unit, price, deposit,
                    packaging_type, price_per_unit_text, is_bio):
    document = Document()
    section = document.sections[0]

    # Set very small margins to maximize content area
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)

    # Set page size (approx. 4x6 inches seems reasonable for price tags)
    section.page_width = Inches(4)
    section.page_height = Inches(6) # Total page height is 6 inches = 432 Pt

    # Set default font for the document (applied to 'Normal' style)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'

    # --- Adaptive Spacing Logic ---
    # These values can be tweaked to ensure everything fits on one page.
    # Start with minimal spacing and increase as needed for visual balance.
    spacing_after_product_name = 1 # Reduced spacing
    spacing_after_varieties = 3 # Pt
    spacing_after_quantity_line = 3 # Pt
    spacing_after_price = 3 # Pt
    spacing_after_deposit = 3 # Pt
    spacing_after_unit_price = 3 # Pt

    # Define RED color
    RED = RGBColor(0xFF, 0x00, 0x00)

    # Determine if prices should be red
    should_be_red = (product_type == 'Aktion')

    # Add Logo
    if os.path.exists(logo_path):
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(logo_path, width=CONSISTENT_LOGO_WIDTH)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Keep tight spacing after logo to ensure content fits on one page.
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)

    # Main content based on department and type
    if department == 'Getränke' and product_type == 'Aktion':
        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_size=36, bold=True, spacing_after=spacing_after_product_name)
        add_tight_paragraph(document, f"{product_name}", font_size=26, bold=True, spacing_after=spacing_after_product_name)

        # Display "Verschiedene Sorten" if has_varieties is true for Aktion products
        if has_varieties:
            add_tight_paragraph(document, "Verschiedene Sorten", font_size=20, spacing_after=spacing_after_varieties)
        elif sub_product_name: # Still keep sub_product_name if varieties is not chosen
            add_tight_paragraph(document, f"{sub_product_name}", font_size=20, spacing_after=spacing_after_varieties)

        je_line_text = f"Je {quantity_per_pack_str}{unit} Flasche"
        add_tight_paragraph(document, je_line_text, font_size=22, spacing_after=spacing_after_quantity_line, apply_highlight=True, underline=True) # Underline JE line

        # Price format with comma and RED color
        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=55, bold=True, spacing_after=spacing_after_price, font_color=RED)
        if deposit > 0:
            add_tight_paragraph(document, f"Zzgl.: {deposit:,.2f}€ Pfand".replace('.', ','), font_size=14, bold=True, spacing_after=spacing_after_deposit, font_color=RED) # Pfand bold
        if price_per_unit_text:
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=18, bold=True, spacing_after=spacing_after_unit_price, font_color=RED) # Price per unit bold
        if packaging_type:
            add_tight_paragraph(document, packaging_type.upper(), font_size=30, bold=True, underline=True, spacing_after=0)

    elif department == 'Obst&Gemüse':
        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_size=36, bold=True, spacing_after=spacing_after_product_name)
        if product_name:
            add_tight_paragraph(document, f"{product_name}", font_size=26, bold=True, spacing_after=spacing_after_product_name)
        if additional_info:
            add_tight_paragraph(document, f"{additional_info}", font_size=20, spacing_after=spacing_after_varieties)

        je_line_text = f"Je {quantity_per_pack_str}{unit} Packung"
        add_tight_paragraph(document, je_line_text, font_size=22, spacing_after=spacing_after_quantity_line, apply_highlight=True, underline=True) # Underline JE line

        # Price format with comma
        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=55, bold=True, spacing_after=spacing_after_price, font_color=RED if should_be_red else None)
        if price_per_unit_text:
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=18, bold=True, spacing_after=0, font_color=RED if should_be_red else None) # Price per unit bold

    else: # Trocken Sortiment or Getränke (Normalpreis)
        # Add one empty line only if it's Trocken Sortiment
        if department == 'Trocken Sortiment':
            # Add one empty paragraph with small spacing to act as an empty line
            empty_paragraph = document.add_paragraph()
            empty_paragraph.paragraph_format.space_before = Pt(0)
            empty_paragraph.paragraph_format.space_after = Pt(12) 

        if manufacturer:
            add_tight_paragraph(document, f"{manufacturer}", font_size=36 + 2, bold=True, spacing_after=spacing_after_product_name) # +2
        if product_name:
            add_tight_paragraph(document, f"{product_name}", font_size=26 + 2, bold=True, spacing_after=spacing_after_product_name) # +2
        if has_varieties:
            add_tight_paragraph(document, "Verschiedene Sorten", font_size=20 + 2, spacing_after=spacing_after_varieties) # +2

        container_word = "Packung"
        if unit == "ml" or unit == "g":
            if "glas" in quantity_per_pack_str.lower():
                container_word = "Glas"
            elif unit == "ml" and "flasche" in quantity_per_pack_str.lower():
                container_word = "Flasche"
            elif unit == "L" and "flasche" in quantity_per_pack_str.lower():
                container_word = "Flasche"
            elif department == 'Getränke':
                container_word = "Flasche"
            else:
                container_word = "Packung"

        je_line_text = f"Je {quantity_per_pack_str}{unit} {container_word}"
        add_tight_paragraph(document, je_line_text, font_size=22 + 2, spacing_after=spacing_after_quantity_line, apply_highlight=True, underline=True) # Underline JE line # +2

        # Price format with comma
        add_tight_paragraph(document, f"{price:,.2f}€".replace('.', ','), font_name='Times New Roman', font_size=55 + 2, bold=True, spacing_after=spacing_after_price, font_color=RED if should_be_red else None) # +2
        if price_per_unit_text:
            if "1kg=" in price_per_unit_text:
                price_per_unit_text = price_per_unit_text
            elif "1L=" in price_per_unit_text or "1l=" in price_per_unit_text:
                price_per_unit_text = price_per_unit_text.replace("1L=", "1L/").replace("1l=", "1L/")
            add_tight_paragraph(document, price_per_unit_text, font_name='Times New Roman', font_size=18 + 2, bold=True, spacing_after=spacing_after_unit_price, font_color=RED if should_be_red else None) # +2
        if department == 'Getränke' and deposit > 0:
            add_tight_paragraph(document, f"Zzgl.: {deposit:,.2f}€ Pfand".replace('.', ','), font_size=14 + 2, bold=True, spacing_after=spacing_after_deposit, font_color=RED if should_be_red else None) # +2
        if department == 'Getränke' and packaging_type:
            add_tight_paragraph(document, packaging_type.upper(), font_size=30 + 2, bold=True, underline=True, spacing_after=0) # +2

    return document

# --- Run the Flask application (for local testing, not used by Elastic Beanstalk) ---
if __name__ == '__main__':
    application.run(debug=True) # CHANGED: app -> application